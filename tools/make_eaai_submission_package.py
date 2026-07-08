from __future__ import annotations

import html
import re
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import PageBreak, Paragraph, Preformatted, SimpleDocTemplate, Spacer, Table, TableStyle


ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "papers/tone_aware_vocos_q1_manuscript.md"
OUT = ROOT / "papers/eaai_submission"

TITLE = (
    "Tone-Aware Fourier Vocoding for Vietnamese Speech Synthesis "
    "Using Low-Frequency Reconstruction and Constant-Q Adversarial Learning"
)
RUNNING_TITLE = "Tone-aware Fourier vocoding for Vietnamese speech synthesis"

ABSTRACT = (
    "Vietnamese text-to-speech synthesis is a practical engineering application of artificial "
    "intelligence in which waveform errors can change lexical meaning because tone is phonemic. "
    "This study contributes an artificial-intelligence training framework for tone-aware neural "
    "vocoding: a Fourier-based Vocos generator is kept unchanged at inference, while training is "
    "augmented with low-frequency weighted spectral reconstruction, Constant-Q Transform "
    "adversarial learning, and a frozen Vietnamese tone-classifier feature loss. The engineering "
    "application is high-quality Vietnamese speech synthesis from compressed MP3 training audio "
    "without corpus-wide WAV conversion. Four vocoders are trained from scratch for one million "
    "steps under matched data, architecture, optimizer, and seed settings. Objective evaluation "
    "uses 20,000 source-disjoint Vietnamese syllable segments derived from forced alignment. "
    "Standard mel validation favors the baseline, showing a reconstruction tradeoff. However, "
    "the best tone-aware system, combining low-frequency spectral weighting and Constant-Q "
    "adversarial learning, improves automatic tone accuracy from 57.48% to 57.73%, reduces "
    "fundamental-frequency error by 6.47%, and reduces voiced/unvoiced error by 2.58% relative "
    "to baseline. The frozen tone-classifier feature loss is a negative ablation. Downstream "
    "integration with a Vietnamese ZipVoice system confirms inference compatibility without "
    "additional architecture cost. These results support a targeted conclusion: tone-aware "
    "training objectives can improve Vietnamese tonal and pitch proxy metrics while preserving "
    "the deployment simplicity of the original vocoder."
)

KEYWORDS = [
    "Vietnamese",
    "Vocoding",
    "Tone",
    "Constant-Q",
    "Speech synthesis",
    "Vocos",
]

HIGHLIGHTS = [
    "Tone-aware Vocos is trained on 3.80M Vietnamese MP3 clips",
    "Low-frequency spectral weighting reduces pitch error by 6.47%",
    "Constant-Q adversarial learning gives the best tone metrics",
    "The tone-classifier perceptual loss is a negative ablation",
    "Tone-aware modules add no inference-time architecture cost",
]

GEN_AI_DECLARATION = (
    "During the preparation of this work, the authors used OpenAI Codex/ChatGPT to assist with "
    "manuscript organization, language editing, formatting, and summarization of experimental "
    "artifacts. After using this tool, the authors reviewed and edited the content as needed and "
    "take full responsibility for the content of the published article."
)

DATA_AVAILABILITY = (
    "The objective evaluation summaries, configuration files, and analysis scripts can be made "
    "available in an anonymized public repository upon acceptance, subject to repository and "
    "license review. The raw MP3 speech corpus cannot be redistributed by the authors because it "
    "contains third-party audio subject to licensing restrictions. The no-overlap evaluation "
    "manifests can be released with anonymized identifiers where legally permitted. The trained "
    "checkpoints can be shared only if the underlying data license allows derived-model release."
)


def extract_body(source: str) -> str:
    start = source.index("## 1. Introduction")
    body = source[start:]
    body = body.replace(
        "low-frequency and harmonic structures associated with F0 and tone realization",
        "low-frequency and harmonic structures associated with fundamental frequency (F0) and tone realization",
        1,
    )
    body = body.replace(
        "rendered through an inverse STFT head",
        "rendered through an inverse short-time Fourier transform (STFT) head",
        1,
    )
    body = body.replace(
        "a CQT tone-focused discriminator that applies adversarial pressure",
        "a Constant-Q Transform (CQT) tone-focused discriminator that applies adversarial pressure",
        1,
    )
    body = re.sub(
        r"## Data and code availability\n.*?\n## Ethical considerations",
        "## Data availability\n\n" + DATA_AVAILABILITY + "\n\n## Ethical considerations",
        body,
        flags=re.S,
    )
    body = re.sub(r"\n## Acknowledgments\n\n`?\[\[ACKNOWLEDGMENTS_PLACEHOLDER\]\]`?\n", "\n", body)
    body = body.replace("/media/huy/data1tb/huy/data_vox/audio", "[anonymized MP3 corpus root]")
    body = body.replace(
        "/data_nvme/ZipVoice/egs/zipvoice/exp/zipvoice_vi_simple_finetune/",
        "[anonymized ZipVoice checkpoint directory]",
    )
    body = body.replace(
        "`/media/huy/data1tb/huy/data_vox/audio`",
        "`[anonymized MP3 corpus root]`",
    )
    body = body.replace(
        "`/data_nvme/ZipVoice/egs/zipvoice/exp/zipvoice_vi_simple_finetune/`",
        "`[anonymized ZipVoice checkpoint directory]`",
    )
    declaration = (
        "## Declaration of generative AI and AI-assisted technologies in the writing process\n\n"
        + GEN_AI_DECLARATION
        + "\n\n"
    )
    body = body.replace("\n## References\n", "\n" + declaration + "## References\n")
    return body


def build_anonymized_markdown() -> str:
    source = SRC.read_text(encoding="utf-8")
    body = extract_body(source)
    return (
        f"# {TITLE}\n\n"
        "## Abstract\n\n"
        f"{ABSTRACT}\n\n"
        f"**Keywords:** {', '.join(KEYWORDS)}\n\n"
        + body
    )


def split_table_row(line: str) -> list[str]:
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [cell.strip() for cell in line.split("|")]


def is_table_sep(line: str) -> bool:
    if not line.strip().startswith("|"):
        return False
    cells = split_table_row(line)
    return all(re.fullmatch(r":?-{3,}:?", cell.strip()) for cell in cells if cell.strip())


def inline_plain(text: str) -> str:
    text = re.sub(r"!\[([^\]]*)\]\([^)]*\)", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1 (\2)", text)
    text = text.replace("**", "").replace("`", "")
    return text.strip()


def add_markdown_to_docx(markdown: str, path: Path, *, title_page: bool = False) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(11)
    for style_name, size in [("Heading 1", 15), ("Heading 2", 13), ("Heading 3", 12)]:
        styles[style_name].font.name = "Times New Roman"
        styles[style_name].font.size = Pt(size)
        styles[style_name].font.bold = True

    lines = markdown.splitlines()
    i = 0
    para: list[str] = []

    def flush_para() -> None:
        nonlocal para
        if para:
            doc.add_paragraph(inline_plain(" ".join(para)))
            para = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        if not stripped:
            flush_para()
            i += 1
            continue
        if stripped.startswith("```"):
            flush_para()
            i += 1
            buf: list[str] = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                buf.append(lines[i])
                i += 1
            if i < len(lines):
                i += 1
            p = doc.add_paragraph()
            run = p.add_run("\n".join(buf))
            run.font.name = "Courier New"
            run.font.size = Pt(8)
            continue
        if stripped.startswith("|") and i + 1 < len(lines) and is_table_sep(lines[i + 1]):
            flush_para()
            table_lines = [line, lines[i + 1]]
            i += 2
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            rows = [split_table_row(row) for row in table_lines if not is_table_sep(row)]
            if rows:
                ncols = max(len(row) for row in rows)
                table = doc.add_table(rows=len(rows), cols=ncols)
                table.style = "Table Grid"
                for r, row in enumerate(rows):
                    row = row + [""] * (ncols - len(row))
                    for c, cell in enumerate(row):
                        table.cell(r, c).text = inline_plain(cell)
                        if r == 0:
                            for paragraph in table.cell(r, c).paragraphs:
                                for run in paragraph.runs:
                                    run.bold = True
            continue
        if stripped.startswith("#"):
            flush_para()
            level = len(stripped) - len(stripped.lstrip("#"))
            text = inline_plain(stripped[level:].strip())
            if level == 1:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(text)
                run.bold = True
                run.font.size = Pt(16)
            else:
                doc.add_heading(text, level=min(level - 1, 3))
            i += 1
            continue
        bullet_match = re.match(r"^[-*+]\s+(.*)$", stripped)
        number_match = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if bullet_match:
            flush_para()
            doc.add_paragraph(inline_plain(bullet_match.group(1)), style="List Bullet")
            i += 1
            continue
        if number_match:
            flush_para()
            doc.add_paragraph(inline_plain(number_match.group(2)), style="List Number")
            i += 1
            continue
        para.append(stripped)
        i += 1
    flush_para()
    doc.save(path)


def add_simple_docx(title: str, lines: list[str], path: Path) -> None:
    markdown = "# " + title + "\n\n" + "\n\n".join(lines)
    add_markdown_to_docx(markdown, path)


def setup_pdf_styles():
    font_dir = Path("/usr/share/fonts/truetype/dejavu")
    pdfmetrics.registerFont(TTFont("DejaVuSans", str(font_dir / "DejaVuSans.ttf")))
    pdfmetrics.registerFont(TTFont("DejaVuSans-Bold", str(font_dir / "DejaVuSans-Bold.ttf")))
    pdfmetrics.registerFont(TTFont("DejaVuSansMono", str(font_dir / "DejaVuSansMono.ttf")))
    styles = getSampleStyleSheet()
    styles.add(
        ParagraphStyle(
            name="DocTitle",
            parent=styles["Title"],
            fontName="DejaVuSans-Bold",
            fontSize=18,
            leading=23,
            alignment=TA_CENTER,
            spaceAfter=14,
        )
    )
    styles.add(
        ParagraphStyle(
            name="H1Custom",
            parent=styles["Heading1"],
            fontName="DejaVuSans-Bold",
            fontSize=15,
            leading=19,
            spaceBefore=14,
            spaceAfter=8,
        )
    )
    styles.add(
        ParagraphStyle(
            name="H2Custom",
            parent=styles["Heading2"],
            fontName="DejaVuSans-Bold",
            fontSize=12.5,
            leading=16,
            spaceBefore=10,
            spaceAfter=6,
        )
    )
    styles.add(
        ParagraphStyle(
            name="BodyCustom",
            parent=styles["BodyText"],
            fontName="DejaVuSans",
            fontSize=9.2,
            leading=12.6,
            spaceAfter=5.2,
            alignment=TA_LEFT,
        )
    )
    styles.add(
        ParagraphStyle(
            name="BulletCustom",
            parent=styles["BodyText"],
            fontName="DejaVuSans",
            fontSize=9.0,
            leading=12.3,
            leftIndent=14,
            firstLineIndent=-10,
            spaceAfter=3.5,
        )
    )
    styles.add(
        ParagraphStyle(
            name="CodeCustom",
            parent=styles["Code"],
            fontName="DejaVuSansMono",
            fontSize=7.2,
            leading=9.2,
            leftIndent=8,
            rightIndent=8,
            backColor=colors.HexColor("#F4F4F4"),
            borderPadding=5,
            spaceBefore=4,
            spaceAfter=6,
        )
    )
    styles.add(
        ParagraphStyle(
            name="TableCell",
            parent=styles["BodyText"],
            fontName="DejaVuSans",
            fontSize=7.1,
            leading=8.7,
            spaceAfter=0,
        )
    )
    styles.add(
        ParagraphStyle(
            name="TableHeader",
            parent=styles["BodyText"],
            fontName="DejaVuSans-Bold",
            fontSize=7.2,
            leading=8.8,
            textColor=colors.white,
            spaceAfter=0,
        )
    )
    return styles


def inline_pdf(text: str) -> str:
    text = re.sub(r"!\[([^\]]*)\]\([^)]*\)", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1 (\2)", text)
    text = html.escape(text.replace("`", ""), quote=False)
    text = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", text)
    return text


def render_pdf(markdown: str, path: Path) -> None:
    styles = setup_pdf_styles()
    story = []
    lines = markdown.splitlines()
    i = 0
    para: list[str] = []

    def flush_para() -> None:
        nonlocal para
        if para:
            story.append(Paragraph(inline_pdf(" ".join(para)), styles["BodyCustom"]))
            para = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        if not stripped:
            flush_para()
            i += 1
            continue
        if stripped.startswith("```"):
            flush_para()
            i += 1
            buf = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                buf.append(lines[i])
                i += 1
            if i < len(lines):
                i += 1
            story.append(Preformatted("\n".join(buf), styles["CodeCustom"]))
            continue
        if stripped.startswith("|") and i + 1 < len(lines) and is_table_sep(lines[i + 1]):
            flush_para()
            buf = [line, lines[i + 1]]
            i += 2
            while i < len(lines) and lines[i].strip().startswith("|"):
                buf.append(lines[i])
                i += 1
            rows = [split_table_row(row) for row in buf if not is_table_sep(row)]
            ncols = max(len(row) for row in rows)
            for row in rows:
                row.extend([""] * (ncols - len(row)))
            char_weights = []
            for c in range(ncols):
                max_len = max(len(re.sub(r"`|\*", "", row[c])) for row in rows)
                char_weights.append(max(8, min(max_len, 28)))
            total = sum(char_weights)
            col_widths = [(A4[0] - 2.4 * cm) * w / total for w in char_weights]
            data = []
            for ri, row in enumerate(rows):
                style_name = "TableHeader" if ri == 0 else "TableCell"
                data.append([Paragraph(inline_pdf(cell), styles[style_name]) for cell in row])
            table = Table(data, colWidths=col_widths, repeatRows=1, hAlign="LEFT")
            table.setStyle(
                TableStyle(
                    [
                        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#3A3A3A")),
                        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                        ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#BDBDBD")),
                        ("VALIGN", (0, 0), (-1, -1), "TOP"),
                        ("LEFTPADDING", (0, 0), (-1, -1), 3.0),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 3.0),
                        ("TOPPADDING", (0, 0), (-1, -1), 3.0),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 3.0),
                        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F7F7F7")]),
                    ]
                )
            )
            story += [Spacer(1, 3), table, Spacer(1, 7)]
            continue
        if stripped.startswith("#"):
            flush_para()
            level = len(stripped) - len(stripped.lstrip("#"))
            text = stripped[level:].strip()
            if level == 1:
                story.append(Paragraph(inline_pdf(text), styles["DocTitle"]))
            elif level == 2:
                if text == "References":
                    story.append(PageBreak())
                story.append(Paragraph(inline_pdf(text), styles["H1Custom"]))
            else:
                story.append(Paragraph(inline_pdf(text), styles["H2Custom"]))
            i += 1
            continue
        bullet_match = re.match(r"^[-*+]\s+(.*)$", stripped)
        number_match = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if bullet_match:
            flush_para()
            story.append(Paragraph(inline_pdf(bullet_match.group(1)), styles["BulletCustom"], bulletText="•"))
            i += 1
            continue
        if number_match:
            flush_para()
            story.append(Paragraph(inline_pdf(number_match.group(2)), styles["BulletCustom"], bulletText=number_match.group(1) + "."))
            i += 1
            continue
        para.append(stripped)
        i += 1
    flush_para()

    doc = SimpleDocTemplate(
        str(path),
        pagesize=A4,
        rightMargin=1.2 * cm,
        leftMargin=1.2 * cm,
        topMargin=1.25 * cm,
        bottomMargin=1.25 * cm,
        title=TITLE,
        author="Anonymous",
    )

    def header_footer(canvas, doc_obj):
        canvas.saveState()
        canvas.setFont("DejaVuSans", 7.5)
        canvas.setFillColor(colors.HexColor("#555555"))
        canvas.drawRightString(A4[0] - 1.2 * cm, 0.65 * cm, f"Page {doc_obj.page}")
        if doc_obj.page > 1:
            canvas.drawString(1.2 * cm, A4[1] - 0.75 * cm, RUNNING_TITLE)
            canvas.line(1.2 * cm, A4[1] - 0.9 * cm, A4[0] - 1.2 * cm, A4[1] - 0.9 * cm)
        canvas.restoreState()

    doc.build(story, onFirstPage=header_footer, onLaterPages=header_footer)


def make_side_files() -> None:
    title_lines = [
        f"**Full title:** {TITLE}",
        f"**Running title:** {RUNNING_TITLE}",
        "**Article type:** Research article",
        "**Authors:** Huy [Surname]^1, [Co-author]^2",
        "**Affiliations:** ^1 [Department, Institution, City, Country]; ^2 [Department, Institution, City, Country]",
        "**Corresponding author:** [Name], [Full postal address], [Email]",
        "**Acknowledgements:** [[ACKNOWLEDGMENTS_PLACEHOLDER]]",
        "**Funding:** This research did not receive any specific grant from funding agencies in the public, commercial, or not-for-profit sectors. [[EDIT_IF_FUNDED]]",
        "**Declaration of competing interest:** The authors declare that they have no known competing financial interests or personal relationships that could have appeared to influence the work reported in this paper. [[EDIT_IF_NEEDED]]",
        "**CRediT author statement:** Huy [Surname]: Conceptualization, Methodology, Software, Data curation, Formal analysis, Investigation, Validation, Writing – original draft, Writing – review and editing. [Co-author]: Supervision, Methodology, Writing – review and editing. [[EDIT_AUTHORS]]",
        "**Data availability:** " + DATA_AVAILABILITY,
    ]
    add_simple_docx("Title page", title_lines, OUT / "EAAI_title_page_author_details.docx")
    (OUT / "EAAI_title_page_author_details.md").write_text("# Title page\n\n" + "\n\n".join(title_lines) + "\n", encoding="utf-8")

    highlights_md = "# Highlights\n\n" + "\n".join(f"- {item}" for item in HIGHLIGHTS) + "\n"
    (OUT / "EAAI_highlights.docx").unlink(missing_ok=True)
    add_markdown_to_docx(highlights_md, OUT / "EAAI_highlights.docx")
    (OUT / "EAAI_highlights.md").write_text(highlights_md, encoding="utf-8")

    cover = [
        "Dear Editor-in-Chief,",
        (
            "We submit the manuscript entitled “"
            + TITLE
            + "” for consideration as a Research Article in Engineering Applications of Artificial Intelligence."
        ),
        (
            "The manuscript addresses a real-world engineering application of artificial intelligence: Vietnamese "
            "speech synthesis, where waveform generation errors can alter lexical tone and therefore word identity. "
            "The artificial-intelligence contribution is a tone-aware neural-vocoder training framework that keeps "
            "the Vocos generator unchanged at inference while adding low-frequency weighted reconstruction and "
            "Constant-Q adversarial learning during training."
        ),
        (
            "The study reports controlled one-million-step ablations trained from scratch on 3.80 million Vietnamese "
            "MP3 clips and validates the best method on a source-disjoint 20,000-segment tone/F0 evaluation. The best "
            "configuration reduces fundamental-frequency error by 6.47% and voiced/unvoiced error by 2.58% relative "
            "to the baseline, while remaining compatible with downstream ZipVoice synthesis."
        ),
        (
            "The manuscript is original, is not under consideration elsewhere, and does not propose a metaphor-based "
            "metaheuristic. All authors have approved the submission. The paper is prepared for double-anonymized "
            "review with a separate title page and anonymized manuscript."
        ),
        "Sincerely,",
        "[Corresponding author name]",
    ]
    add_simple_docx("Cover letter", cover, OUT / "EAAI_cover_letter.docx")
    (OUT / "EAAI_cover_letter.md").write_text("# Cover letter\n\n" + "\n\n".join(cover) + "\n", encoding="utf-8")

    declarations = [
        "**Declaration of competing interest:** The authors declare that they have no known competing financial interests or personal relationships that could have appeared to influence the work reported in this paper. [[EDIT_IF_NEEDED]]",
        "**Funding:** This research did not receive any specific grant from funding agencies in the public, commercial, or not-for-profit sectors. [[EDIT_IF_FUNDED]]",
        "**Declaration of generative AI and AI-assisted technologies in the writing process:** " + GEN_AI_DECLARATION,
        "**Data availability:** " + DATA_AVAILABILITY,
        "**Ethics statement for listening tests:** [[ADD_AFTER_MOS_OR_TONE_ID_STUDY: listener consent, recruitment, screening, compensation, and institutional review status]]",
    ]
    add_simple_docx("Declarations", declarations, OUT / "EAAI_declarations.docx")
    (OUT / "EAAI_declarations.md").write_text("# Declarations\n\n" + "\n\n".join(declarations) + "\n", encoding="utf-8")

    checklist = [
        "# EAAI submission checklist",
        "",
        "Source checked from the official ScienceDirect/Elsevier Guide for Authors on 2026-07-06.",
        "",
        "## Completed in this package",
        "",
        "- Single-column anonymized manuscript created as DOCX and PDF.",
        "- Separate title page with author placeholders created.",
        "- Separate highlights file created; all five highlights are <=85 characters.",
        "- Abstract shortened to 204 words and explicitly states AI contribution plus engineering application.",
        "- Undefined acronym risk reduced in title and abstract.",
        "- Declaration of generative AI use added before references.",
        "- Cover letter drafted for Engineering Applications of Artificial Intelligence.",
        "- Declarations file drafted.",
        "",
        "## Must be fixed before real submission",
        "",
        "- Replace all `[[...]]` placeholders.",
        "- Fill final author names, affiliations, full corresponding-author address, and email.",
        "- Complete Elsevier Declaration of Interest tool and upload the generated DOCX if required.",
        "- Decide whether to keep subjective placeholders. Do not submit with placeholder MOS/MUSHRA/tone-ID values.",
        "- If no human listening test is completed, remove placeholder tables and frame them as future work.",
        "- Strongly consider adding public-dataset validation because EAAI states submitted papers should be validated using public datasets for replicability.",
        "- Review whether raw data/checkpoint licenses permit public release or derived-model sharing.",
        "- Optional: prepare graphical abstract without generative-AI-created artwork.",
        "",
        "## Desk-rejection risk notes",
        "",
        "- EAAI says the abstract must clearly specify the AI contribution and the engineering application; this package addresses that directly.",
        "- EAAI follows double-anonymized review; do not include author names or acknowledgements in the anonymized manuscript.",
        "- EAAI expects public-data validation for replicability. The current corpus is local/private, so this is the main strategic risk.",
        "- Placeholders are not acceptable in final submission files.",
    ]
    (OUT / "EAAI_submission_checklist.md").write_text("\n".join(checklist) + "\n", encoding="utf-8")
    add_markdown_to_docx("\n".join(checklist), OUT / "EAAI_submission_checklist.docx")


def main() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    manuscript = build_anonymized_markdown()
    md_path = OUT / "EAAI_anonymized_manuscript_with_placeholders.md"
    docx_path = OUT / "EAAI_anonymized_manuscript_with_placeholders.docx"
    pdf_path = OUT / "EAAI_anonymized_manuscript_with_placeholders.pdf"
    md_path.write_text(manuscript, encoding="utf-8")
    add_markdown_to_docx(manuscript, docx_path)
    render_pdf(manuscript, pdf_path)
    make_side_files()

    zip_path = OUT / "EAAI_submission_package_precheck.zip"
    with zipfile.ZipFile(zip_path, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        for path in sorted(OUT.iterdir()):
            if path == zip_path or path.is_dir():
                continue
            zf.write(path, arcname=path.name)
    print(f"Wrote {OUT}")
    print(f"Wrote {zip_path}")


if __name__ == "__main__":
    main()
